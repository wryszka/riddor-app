"""Generate company-format Risk Assessment documents (PDF + DOCX).

The form is a landscape table with one row per hazard, matching the columns
described in the requirements: Hazard + Potential Harm | Pre-control L/C/Rating
| Control Measures | Post-control L/C/Rating | Result.
"""

import io
from datetime import date

from fpdf import FPDF
from fpdf.enums import VAlign
from fpdf.fonts import FontFace


# Shared with the PDF font (Latin-1) — replace common unicode chars.
_REPLACEMENTS = {
    "‐": "-", "‑": "-", "‒": "-", "–": "-", "—": "-", "―": "-", "−": "-",
    "•": "-", "‣": "-", "◦": "-", "·": "-",
    "‘": "'", "’": "'", "‚": ",", "‛": "'",
    "“": '"', "”": '"', "„": '"', "‟": '"',
    "≤": "<=", "≥": ">=", "≠": "!=", "≈": "~", "×": "x", "÷": "/",
    "°": " deg", "μ": "u", "µ": "u", "→": "->", "←": "<-", "…": "...",
    " ": " ", "®": "(R)", "©": "(c)", "™": "(TM)",
}


def _safe(text) -> str:
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    for old, new in _REPLACEMENTS.items():
        text = text.replace(old, new)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _result_rgb(result: str) -> tuple[int, int, int]:
    r = (result or "").upper()
    if "VERY HIGH" in r:
        return (153, 27, 27)
    if "HIGH" in r:
        return (220, 38, 38)
    if "MEDIUM" in r:
        return (217, 119, 6)
    return (22, 163, 74)  # LOW


def _harm_text(row: dict) -> str:
    lines = [row.get("hazard", "")]
    harm = row.get("potential_harm") or []
    if harm:
        lines.append("")
        lines.append("Potential harm:")
        lines.extend(f"- {h}" for h in harm)
    persons = row.get("persons_at_risk") or []
    if persons:
        lines.append("")
        lines.append("At risk: " + ", ".join(persons))
    return "\n".join(lines)


def _controls_text(row: dict) -> str:
    controls = row.get("control_measures") or []
    return "\n".join(f"- {c}" for c in controls)


# ── PDF ───────────────────────────────────────────────────────────────

class _RAPDF(FPDF):
    title_text = "RISK ASSESSMENT"
    task_text = ""

    def header(self):
        self.set_font("Helvetica", "B", 14)
        self.set_text_color(15, 23, 42)
        self.cell(0, 8, _safe(self.title_text), ln=True, align="C")
        if self.task_text:
            self.set_font("Helvetica", "", 10)
            self.set_text_color(71, 85, 105)
            self.cell(0, 5, _safe(self.task_text), ln=True, align="C")
        self.set_draw_color(37, 99, 235)
        self.set_line_width(0.6)
        self.line(self.l_margin, self.get_y() + 1, self.w - self.r_margin, self.get_y() + 1)
        self.ln(4)

    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", "I", 7)
        self.set_text_color(150, 150, 150)
        self.cell(
            0, 8,
            f"AI-generated risk assessment - review and approve before use | Page {self.page_no()}",
            align="C",
        )


def build_ra_pdf(ra: dict) -> bytes:
    """Build a landscape Risk Assessment PDF, one row per hazard."""
    pdf = _RAPDF(orientation="L", unit="mm", format="A4")
    pdf.title_text = "RISK ASSESSMENT"
    pdf.task_text = ra.get("task_title") or ""
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.add_page()

    # Metadata strip
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(71, 85, 105)
    meta = f"Date: {date.today().isoformat()}"
    activities = ra.get("activities") or []
    if activities:
        meta += "    Activities: " + ", ".join(activities)
    pdf.multi_cell(0, 5, _safe(meta))
    pdf.ln(2)

    rows = ra.get("hazards") or []

    # Column widths (sum ~ 269mm usable in landscape A4 with 14mm margins)
    col_widths = (58, 24, 10, 10, 13, 92, 10, 10, 13, 19)
    headings = (
        "Hazard / Potential Harm", "Persons\nat Risk",
        "L", "C", "Risk", "Control Measures",
        "L", "C", "Risk", "Result",
    )

    head_style = FontFace(emphasis="BOLD", color=(255, 255, 255), fill_color=(37, 99, 235))
    grey = (241, 245, 249)

    with pdf.table(
        col_widths=col_widths,
        text_align=("LEFT", "CENTER", "CENTER", "CENTER", "CENTER", "LEFT", "CENTER", "CENTER", "CENTER", "CENTER"),
        line_height=5,
        v_align=VAlign.T,
        headings_style=head_style,
        first_row_as_headings=True,
        borders_layout="ALL",
    ) as table:
        hrow = table.row()
        for h in headings:
            hrow.cell(_safe(h))

        for i, r in enumerate(rows):
            zebra = grey if i % 2 else (255, 255, 255)
            res_rgb = _result_rgb(r.get("result"))
            row = table.row()
            row.cell(_safe(_harm_text(r)), style=FontFace(fill_color=zebra))
            row.cell(_safe(", ".join(r.get("persons_at_risk") or [])), style=FontFace(fill_color=zebra))
            row.cell(str(r.get("pre_likelihood", "")), style=FontFace(fill_color=zebra))
            row.cell(str(r.get("pre_consequence", "")), style=FontFace(fill_color=zebra))
            row.cell(str(r.get("pre_rating", "")), style=FontFace(emphasis="BOLD", fill_color=zebra))
            row.cell(_safe(_controls_text(r)), style=FontFace(fill_color=zebra))
            row.cell(str(r.get("post_likelihood", "")), style=FontFace(fill_color=zebra))
            row.cell(str(r.get("post_consequence", "")), style=FontFace(fill_color=zebra))
            row.cell(str(r.get("post_rating", "")), style=FontFace(emphasis="BOLD", fill_color=zebra))
            row.cell(_safe(r.get("result", "")), style=FontFace(emphasis="BOLD", color=(255, 255, 255), fill_color=res_rgb))

    out = pdf.output()
    return bytes(out) if isinstance(out, (bytes, bytearray)) else out.encode("latin-1")


# ── DOCX ──────────────────────────────────────────────────────────────

def _shade_cell(cell, rgb_hex: str):
    """Apply background shading to a docx table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex)
    tcPr.append(shd)


def build_ra_docx(ra: dict) -> bytes:
    """Build a company-format Risk Assessment as a .docx with a hazard table."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    doc = Document()
    section = doc.sections[0]
    # Landscape
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for m in ("left_margin", "right_margin"):
        setattr(section, m, Cm(1.3))

    title = doc.add_heading("Risk Assessment", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if ra.get("task_title"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(ra["task_title"])
        run.bold = True
        run.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.add_run(f"Date: {date.today().isoformat()}").font.size = Pt(9)
    if ra.get("activities"):
        meta.add_run("    Activities: " + ", ".join(ra["activities"])).font.size = Pt(9)

    headings = [
        "Hazard / Potential Harm", "Persons at Risk",
        "Pre L", "Pre C", "Pre Risk", "Control Measures",
        "Post L", "Post C", "Post Risk", "Result",
    ]
    rows = ra.get("hazards") or []
    table = doc.add_table(rows=1, cols=len(headings))
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, h in enumerate(headings):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hdr[i], "2563EB")

    result_fill = {
        "LOW": "16A34A", "MEDIUM": "D97706", "HIGH": "DC2626", "VERY HIGH": "991B1B",
    }

    for r in rows:
        cells = table.add_row().cells
        cells[0].text = _harm_text(r)
        cells[1].text = ", ".join(r.get("persons_at_risk") or [])
        cells[2].text = str(r.get("pre_likelihood", ""))
        cells[3].text = str(r.get("pre_consequence", ""))
        cells[4].text = str(r.get("pre_rating", ""))
        cells[5].text = _controls_text(r)
        cells[6].text = str(r.get("post_likelihood", ""))
        cells[7].text = str(r.get("post_consequence", ""))
        cells[8].text = str(r.get("post_rating", ""))
        result = r.get("result", "")
        cells[9].text = ""
        rrun = cells[9].paragraphs[0].add_run(result)
        rrun.bold = True
        rrun.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cells[9], result_fill.get(result.upper(), "16A34A"))
        # Shrink body font for density
        for c in cells:
            for para in c.paragraphs:
                for run in para.runs:
                    if not run.font.size:
                        run.font.size = Pt(8)

    note = doc.add_paragraph()
    nr = note.add_run("AI-generated risk assessment — review and approve before use.")
    nr.italic = True
    nr.font.size = Pt(8)
    nr.font.color.rgb = RGBColor(0x96, 0x96, 0x96)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
