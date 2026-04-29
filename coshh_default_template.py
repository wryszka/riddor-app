"""Built-in COSHH docx template — generated at runtime via python-docx.

This produces a single-page COSHH risk assessment matching the structure
of common UK H&S templates, with Jinja-style placeholders that
``coshh_docx.fill_template`` will populate.
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _shade_cell(cell, hex_color: str):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_run_font(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _label_value_row(table, label: str, placeholder: str):
    """Add a 2-column row: bold label | placeholder cell."""
    row = table.add_row()
    cells = row.cells
    cells[0].text = ""
    p = cells[0].paragraphs[0]
    r = p.add_run(label)
    _set_run_font(r, size=10, bold=True)
    _shade_cell(cells[0], "F1F5F9")

    cells[1].text = ""
    p2 = cells[1].paragraphs[0]
    r2 = p2.add_run("{{" + placeholder + "}}")
    _set_run_font(r2, size=10)


def _full_width_row(table, label: str, placeholder: str):
    """Add a row spanning both columns: bold label, then a row with the placeholder."""
    # Label row (merged)
    row = table.add_row()
    cells = row.cells
    cells[0].merge(cells[1])
    cells[0].text = ""
    p = cells[0].paragraphs[0]
    r = p.add_run(label)
    _set_run_font(r, size=10, bold=True, color="2563EB")

    # Value row (merged)
    row2 = table.add_row()
    c2 = row2.cells
    c2[0].merge(c2[1])
    c2[0].text = ""
    p2 = c2[0].paragraphs[0]
    r2 = p2.add_run("{{" + placeholder + "}}")
    _set_run_font(r2, size=10)


def build_default_template() -> bytes:
    """Build a COSHH risk assessment template with placeholders."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("CONDENSED COSHH ASSESSMENT")
    _set_run_font(tr, size=16, bold=True, color="0F172A")
    doc.add_paragraph()

    # Header table — 2 columns: label / value
    header_table = doc.add_table(rows=0, cols=2)
    header_table.style = "Table Grid"
    header_table.autofit = False
    # Set column widths
    for row in header_table.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.3)

    _label_value_row(header_table, "Product Name", "product_name")
    _label_value_row(header_table, "Use of Product", "use_of_product")
    _label_value_row(header_table, "Product Code", "product_code")
    _label_value_row(header_table, "Persons at Risk", "persons_at_risk")
    _label_value_row(header_table, "Form and Colour", "form_and_colour")
    _label_value_row(header_table, "Risk Rating (after controls)", "risk_rating")

    doc.add_paragraph()

    # Sections table — full width per section
    sections_table = doc.add_table(rows=0, cols=2)
    sections_table.style = "Table Grid"

    _full_width_row(sections_table, "Method of Application", "method_of_application")
    _full_width_row(sections_table, "CLP Hazard Information", "clp_hazard_information")
    _full_width_row(sections_table, "Routes of Entry", "routes_of_entry")

    # PPE block — sub-rows
    row_ppe_header = sections_table.add_row()
    cells = row_ppe_header.cells
    cells[0].merge(cells[1])
    cells[0].text = ""
    p = cells[0].paragraphs[0]
    r = p.add_run("PPE Required")
    _set_run_font(r, size=10, bold=True, color="2563EB")

    for label, placeholder in [
        ("Hand Protection (EN 374/420)", "ppe_hands"),
        ("Eye Protection (EN 166)", "ppe_eyes"),
        ("Respiratory Protection", "ppe_respiratory"),
        ("Protective Clothing", "ppe_clothing"),
        ("Protective Footwear (EN ISO 20345 SR)", "ppe_footwear"),
    ]:
        _label_value_row(sections_table, label, placeholder)

    # First aid block
    row_fa_header = sections_table.add_row()
    cells = row_fa_header.cells
    cells[0].merge(cells[1])
    cells[0].text = ""
    p = cells[0].paragraphs[0]
    r = p.add_run("First Aid Treatment")
    _set_run_font(r, size=10, bold=True, color="2563EB")

    for label, placeholder in [
        ("Skin contact", "first_aid_skin"),
        ("Eye contact", "first_aid_eye"),
        ("Inhalation", "first_aid_inhalation"),
        ("Ingestion", "first_aid_ingestion"),
    ]:
        _label_value_row(sections_table, label, placeholder)

    _full_width_row(sections_table, "Spillage Procedure", "spillage_procedure")
    _full_width_row(sections_table, "Handling and Storage", "handling_and_storage")
    _full_width_row(sections_table, "General Precautions", "general_precautions")
    _full_width_row(sections_table, "Disposal", "disposal")

    # Acceptance & review (signature block — left blank for the assessor)
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Acceptance and Review")
    _set_run_font(r, size=11, bold=True, color="0F172A")
    p2 = doc.add_paragraph(
        "I confirm that the above detailed COSHH risk assessment appropriately considers the "
        "substance to be used and the associated hazards. I am satisfied that all hazards have "
        "been sufficiently identified and that the recorded control measures will reduce the "
        "risks to as low a level as reasonably practicable."
    )
    for run in p2.runs:
        _set_run_font(run, size=9)

    sig_table = doc.add_table(rows=2, cols=4)
    sig_table.style = "Table Grid"
    headers = ["Date", "Name", "Job Title", "Signature"]
    for i, h in enumerate(headers):
        cell = sig_table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        _set_run_font(r, size=10, bold=True)
        _shade_cell(cell, "F1F5F9")

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("AI-generated COSHH assessment — verify against original SDS before use")
    _set_run_font(fr, size=8, color="6B7280")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
