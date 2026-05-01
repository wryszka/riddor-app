"""RIDDOR + COSHH — NiceGUI application."""

import os
import io
import json
import uuid
from datetime import date, datetime, timedelta

from nicegui import ui, events

from riddor_data import (
    INCIDENT_TYPE_CONFIG,
    STATUS_CONFIG,
    MOCK_INCIDENTS,
    MOCK_ACTIONS,
    get_days_remaining,
)


# ── In-memory state (single-tenant demo) ─────────────────────────────

incidents: dict = {i["id"]: i for i in MOCK_INCIDENTS}
actions: list = list(MOCK_ACTIONS)

# COSHH state
sds_doc: dict | None = None       # {"name", "text", "structured"}
selected_template_key: str = "_builtin_"
coshh_chat_messages: list = []
ai_chat_messages: list = []


def reset_demo():
    global incidents, actions, sds_doc, coshh_chat_messages, ai_chat_messages
    incidents = {i["id"]: i for i in MOCK_INCIDENTS}
    actions = list(MOCK_ACTIONS)
    sds_doc = None
    coshh_chat_messages = []
    ai_chat_messages = []


# ── Theme & shared CSS ───────────────────────────────────────────────

SHARED_HEAD = """
<style>
  body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; }
  .badge { display: inline-block; padding: 2px 10px; border-radius: 9999px; font-size: 12px; font-weight: 600; line-height: 20px; white-space: nowrap; }
  .badge-red { background: #dc2626; color: #fff; }
  .badge-orange { background: #ea580c; color: #fff; }
  .badge-amber { background: #d97706; color: #fff; }
  .badge-blue { background: #2563eb; color: #fff; }
  .badge-purple { background: #7c3aed; color: #fff; }
  .badge-green { background: #16a34a; color: #fff; }
  .badge-gray { background: #6b7280; color: #fff; }
  .badge-dark { background: #111827; color: #fff; }
  .urgency-overdue { color: #dc2626; font-weight: 700; }
  .urgency-soon { color: #d97706; font-weight: 600; }
  .urgency-ok { color: #16a34a; }
  .hero-card { transition: transform .15s ease-out, box-shadow .15s ease-out; cursor: pointer; }
  .hero-card:hover { transform: translateY(-3px); box-shadow: 0 12px 40px rgba(0,0,0,0.12); }
  .nicegui-content { padding: 0 !important; }
  .doc-card { background: #fff; border: 1px solid #e2e8f0; border-radius: 12px; padding: 24px; }
  .field-label { font-weight: 700; color: #0f172a; min-width: 200px; display: inline-block; }
  .section-title { font-weight: 700; color: #2563eb; margin-top: 16px; margin-bottom: 4px; }
</style>
"""


def _inject_styles():
    """Inject shared CSS — must be called inside a page function (NiceGUI 3.x rule)."""
    ui.add_head_html(SHARED_HEAD)


TYPE_BADGE = {
    "death": "badge-dark",
    "specified_injury": "badge-red",
    "over_7_day": "badge-orange",
    "non_worker_hospital": "badge-purple",
    "occupational_disease": "badge-blue",
    "dangerous_occurrence": "badge-amber",
    "not_reportable": "badge-gray",
}
STATUS_BADGE = {
    "open": "badge-blue",
    "investigating": "badge-amber",
    "pending_report": "badge-orange",
    "submitted": "badge-green",
    "closed": "badge-gray",
}


# ── Layout shell ─────────────────────────────────────────────────────

def app_header(active: str = ""):
    """Top app bar with brand and section links."""
    _inject_styles()
    with ui.header().classes("bg-white text-slate-900 shadow-sm").style("border-bottom: 1px solid #e2e8f0"):
        with ui.row().classes("w-full items-center max-w-screen-2xl mx-auto px-4 py-2"):
            with ui.row().classes("items-center cursor-pointer").on("click", lambda: ui.navigate.to("/")):
                ui.icon("shield", size="md").classes("text-blue-600")
                ui.label("H&S Hub").classes("text-lg font-bold")
            ui.space()
            with ui.row().classes("items-center gap-1"):
                _nav_link("Home", "/", active == "home")
                _nav_link("RIDDOR", "/riddor", active == "riddor")
                _nav_link("COSHH", "/coshh", active == "coshh")
                ui.button("🔄 Reset Demo", on_click=lambda: (reset_demo(), ui.notify("Demo reset", type="positive"), ui.navigate.to("/"))).props("flat dense")


def _nav_link(label: str, target: str, is_active: bool):
    classes = "px-3 py-1.5 rounded text-sm font-semibold "
    if is_active:
        classes += "bg-blue-50 text-blue-700"
    else:
        classes += "text-slate-700 hover:bg-slate-100"
    ui.link(label, target).classes(classes).style("text-decoration: none")


def page_container():
    """Returns a centered max-width container for page content."""
    return ui.column().classes("max-w-screen-2xl mx-auto w-full p-6 gap-6")


def section_header(icon: str, title: str, subtitle: str | None = None):
    with ui.row().classes("items-center gap-3 w-full"):
        ui.icon(icon, size="lg").classes("text-blue-600")
        with ui.column().classes("gap-0"):
            ui.label(title).classes("text-2xl font-bold")
            if subtitle:
                ui.label(subtitle).classes("text-sm text-slate-500")


# ── RIDDOR section ───────────────────────────────────────────────────

def riddor_subnav(active: str):
    with ui.row().classes("gap-2 w-full bg-slate-50 px-2 py-1 rounded-lg"):
        for label, target, key in [
            ("📊 Case Monitor", "/riddor", "monitor"),
            ("📝 Report Incident", "/riddor/report", "report"),
            ("📁 Past Reports", "/riddor/archive", "archive"),
            ("💬 AI Assistant", "/riddor/assistant", "assistant"),
        ]:
            cls = "px-3 py-2 rounded text-sm font-semibold "
            cls += "bg-blue-600 text-white" if key == active else "text-slate-700 hover:bg-white"
            ui.link(label, target).classes(cls).style("text-decoration: none")


# ── Landing ──────────────────────────────────────────────────────────

@ui.page("/")
def landing_page():
    app_header(active="home")
    with page_container():
        with ui.column().classes("items-center text-center w-full py-8 gap-2"):
            ui.label("Health & Safety Hub").classes("text-4xl font-bold")
            ui.label("AI-powered RIDDOR decision support and COSHH chemical assessments").classes("text-lg text-slate-500")

        with ui.grid(columns=2).classes("w-full gap-6 mt-4"):
            # RIDDOR card
            with ui.card().tight().classes("hero-card").on("click", lambda: ui.navigate.to("/riddor")):
                with ui.column().classes("p-8 gap-3"):
                    ui.icon("shield", size="3rem").classes("text-blue-600")
                    ui.label("RIDDOR").classes("text-2xl font-bold")
                    ui.label("Reporting of Injuries, Diseases and Dangerous Occurrences Regulations 2013").classes("text-slate-600")
                    with ui.column().classes("gap-1 mt-3"):
                        for item in [
                            "Track open cases and deadlines",
                            "Report new incidents with AI classification",
                            "Search and analyse past reports",
                            "Ask the AI assistant questions",
                        ]:
                            with ui.row().classes("items-center gap-2 text-sm text-slate-700"):
                                ui.icon("check_circle", size="sm").classes("text-green-600")
                                ui.label(item)

            # COSHH card
            with ui.card().tight().classes("hero-card").on("click", lambda: ui.navigate.to("/coshh")):
                with ui.column().classes("p-8 gap-3"):
                    ui.icon("science", size="3rem").classes("text-purple-600")
                    ui.label("COSHH").classes("text-2xl font-bold")
                    ui.label("Control of Substances Hazardous to Health Regulations 2002").classes("text-slate-600")
                    with ui.column().classes("gap-1 mt-3"):
                        for item in [
                            "Upload chemical Safety Data Sheets",
                            "Get instant condensed assessments",
                            "Fill your own COSHH template (.docx)",
                            "Ask questions about each substance",
                        ]:
                            with ui.row().classes("items-center gap-2 text-sm text-slate-700"):
                                ui.icon("check_circle", size="sm").classes("text-green-600")
                                ui.label(item)


# ── RIDDOR: Case Monitor (Dashboard) ─────────────────────────────────

@ui.page("/riddor")
def riddor_dashboard():
    app_header(active="riddor")
    with page_container():
        section_header("dashboard", "Case Monitor", "Track and manage open RIDDOR cases")
        riddor_subnav("monitor")

        all_incidents = list(incidents.values())
        today = date.today().isoformat()
        open_statuses = {"open", "investigating", "pending_report"}
        open_cases = [i for i in all_incidents if i["status"] in open_statuses]
        overdue = [i for i in open_cases if i.get("reporting_deadline") and i["reporting_deadline"] < today]
        approaching = [
            i for i in open_cases
            if i.get("reporting_deadline")
            and i["reporting_deadline"] >= today
            and i["reporting_deadline"] <= (date.today() + timedelta(days=3)).isoformat()
        ]
        submitted = [i for i in all_incidents if i["status"] == "submitted"]

        # Metrics
        with ui.grid(columns=4).classes("w-full gap-4"):
            _metric_card("Open Cases", len(open_cases), "folder_open", "blue")
            _metric_card("Approaching Deadline", len(approaching), "schedule", "amber")
            _metric_card("Overdue", len(overdue), "warning", "red")
            _metric_card("Submitted to HSE", len(submitted), "check_circle", "green")

        # Cases
        if not open_cases:
            with ui.card().classes("w-full p-8 items-center"):
                ui.icon("check_circle", size="3rem").classes("text-green-500")
                ui.label("No open cases — all RIDDOR cases have been resolved").classes("text-slate-500")
        else:
            ui.label(f"Open Cases ({len(open_cases)})").classes("text-lg font-semibold mt-2")
            open_cases.sort(key=lambda i: i.get("reporting_deadline") or "9999-12-31")
            for inc in open_cases:
                _render_case_card(inc)


def _metric_card(label: str, value, icon: str, color: str):
    color_map = {
        "blue": ("text-blue-600", "bg-blue-50"),
        "amber": ("text-amber-600", "bg-amber-50"),
        "red": ("text-red-600", "bg-red-50"),
        "green": ("text-green-600", "bg-green-50"),
    }
    text_cls, bg_cls = color_map.get(color, ("text-slate-600", "bg-slate-50"))
    with ui.card().classes("w-full p-5"):
        with ui.row().classes("items-center justify-between w-full"):
            with ui.column().classes("gap-0"):
                ui.label(label).classes("text-sm text-slate-500 font-medium")
                ui.label(str(value)).classes("text-3xl font-bold")
            with ui.element("div").classes(f"p-3 rounded-xl {bg_cls}"):
                ui.icon(icon, size="md").classes(text_cls)


def _render_case_card(inc: dict):
    days = get_days_remaining(inc.get("reporting_deadline"))
    border = ""
    if days is not None and days < 0:
        border = "border-l-4 border-red-500"
    elif days is not None and days <= 3:
        border = "border-l-4 border-amber-500"

    with ui.card().classes(f"w-full p-4 {border}"):
        with ui.row().classes("items-center w-full gap-4 flex-wrap"):
            with ui.column().classes("gap-0"):
                ui.label(inc["reference"]).classes("font-mono font-semibold")
                ui.label(inc["incident_date"]).classes("text-xs text-slate-500")
            ui.html(_type_badge(inc["incident_type"]))
            with ui.column().classes("gap-0"):
                ui.label(inc["person_name"]).classes("font-semibold")
                ui.label(f"{inc['department']} · {inc['person_type'].replace('_', '-')}").classes("text-xs text-slate-500")
            ui.space()
            ui.html(_status_badge(inc["status"]))
            ui.html(_deadline_html(inc.get("reporting_deadline"), inc["status"]))

        if inc.get("incident_type") == "over_7_day" and inc.get("absence_days") is not None:
            if inc["absence_days"] > 7:
                ui.label(f"📅 Absence: {inc['absence_days']} days — exceeds 7-day RIDDOR threshold").classes("text-amber-700 mt-2 text-sm font-medium")
            else:
                ui.label(f"📅 Absence: {inc['absence_days']} days — monitoring").classes("text-slate-600 mt-2 text-sm")

        with ui.expansion("View details").classes("w-full"):
            ui.label("Description").classes("font-semibold mt-2")
            ui.label(inc["description"]).classes("text-slate-700")
            if inc.get("injury_details"):
                ui.label("Injury").classes("font-semibold mt-2")
                ui.label(inc["injury_details"]).classes("text-slate-700")
            if inc.get("ai_reasoning"):
                ui.label("AI Assessment").classes("font-semibold mt-2")
                ui.label(inc["ai_reasoning"]).classes("text-slate-700")

            inc_actions = [a for a in actions if a["incident_id"] == inc["id"]]
            if inc_actions:
                ui.label("Timeline").classes("font-semibold mt-3")
                for a in inc_actions:
                    ui.label(f"🔹 {a['timestamp'][:10]} — {a['description']} ({a['performed_by']})").classes("text-xs text-slate-600")

            with ui.row().classes("gap-2 mt-3"):
                ui.button("Mark as Submitted", on_click=lambda i=inc: _set_status(i["id"], "submitted")).props("color=positive")
                ui.button("Mark as Investigating", on_click=lambda i=inc: _set_status(i["id"], "investigating")).props("color=warning")
                ui.button("Close Case", on_click=lambda i=inc: _set_status(i["id"], "closed")).props("flat")


def _set_status(case_id: str, new_status: str):
    if case_id in incidents:
        incidents[case_id]["status"] = new_status
        if new_status == "submitted":
            incidents[case_id]["submitted_at"] = date.today().isoformat()
    ui.notify(f"Case marked as {new_status}", type="positive")
    ui.navigate.reload()


def _type_badge(t: str) -> str:
    cfg = INCIDENT_TYPE_CONFIG.get(t, {"label": t})
    cls = TYPE_BADGE.get(t, "badge-gray")
    return f'<span class="badge {cls}">{cfg["label"]}</span>'


def _status_badge(s: str) -> str:
    cfg = STATUS_CONFIG.get(s, {"label": s})
    cls = STATUS_BADGE.get(s, "badge-gray")
    return f'<span class="badge {cls}">{cfg["label"]}</span>'


def _deadline_html(deadline, status: str) -> str:
    if not deadline or status in ("submitted", "closed"):
        return '<span class="text-slate-400">—</span>'
    days = get_days_remaining(deadline)
    if days is None:
        return f'<span>{deadline}</span>'
    if days < 0:
        return f'<span class="urgency-overdue">{deadline}<br/><small>{abs(days)}d overdue</small></span>'
    if days <= 3:
        return f'<span class="urgency-soon">{deadline}<br/><small>{days}d left</small></span>'
    return f'<span class="urgency-ok">{deadline}<br/><small>{days}d left</small></span>'


# ── RIDDOR: Report Incident ──────────────────────────────────────────

@ui.page("/riddor/report")
def report_page():
    app_header(active="riddor")
    with page_container():
        section_header("note_add", "Report an Incident", "Describe an incident for AI-powered RIDDOR classification")
        riddor_subnav("report")

        # State for this page
        page_state = {"step": "input", "classification": None}

        container = ui.column().classes("w-full gap-4")

        def render():
            container.clear()
            with container:
                if page_state["step"] == "input":
                    _render_report_input(page_state, render)
                elif page_state["step"] == "result":
                    _render_report_result(page_state, render)
                elif page_state["step"] == "saved":
                    _render_report_saved(page_state)

        render()


def _render_report_input(state: dict, rerender):
    with ui.card().classes("w-full p-6"):
        ui.label("Describe the incident").classes("text-lg font-semibold mb-2")

        with ui.grid(columns=2).classes("w-full gap-4"):
            with ui.column().classes("gap-3"):
                description = ui.textarea(
                    "What happened?",
                    placeholder="E.g., An employee slipped on a wet floor in the warehouse and broke their wrist..."
                ).props("rows=8 outlined").classes("w-full")
                injury = ui.textarea(
                    "Injury details",
                    placeholder="Type and severity of injury, treatment given..."
                ).props("rows=3 outlined").classes("w-full")

            with ui.column().classes("gap-3"):
                incident_date_input = ui.input("Incident date", value=date.today().isoformat()).props("type=date outlined").classes("w-full")
                person_type = ui.select(["worker", "non_worker"], value="worker", label="Person type").props("outlined").classes("w-full")
                person_name = ui.input("Injured person's name").props("outlined").classes("w-full")
                department = ui.input("Department").props("outlined").classes("w-full")
                location = ui.input("Location").props("outlined").classes("w-full")
                reporter = ui.input("Your name (reporter)").props("outlined").classes("w-full")

        async def classify_now():
            text = description.value or ""
            if not text.strip():
                ui.notify("Please describe the incident", type="warning")
                return
            full = "\n".join(filter(None, [
                text,
                f"Person involved: {person_name.value}" if person_name.value else None,
                f"Person type: {person_type.value}",
                f"Department: {department.value}" if department.value else None,
                f"Location: {location.value}" if location.value else None,
                f"Injury details: {injury.value}" if injury.value else None,
            ]))
            n = ui.notification(message="AI is classifying the incident...", spinner=True, timeout=None)
            try:
                from riddor_ai import classify_incident
                result = classify_incident(full)
                state["classification"] = result
                state["form"] = {
                    "incident_date": incident_date_input.value,
                    "person_type": person_type.value,
                    "person_name": person_name.value,
                    "department": department.value,
                    "location": location.value,
                    "description": text,
                    "injury_details": injury.value,
                    "reporter_name": reporter.value,
                }
                state["step"] = "result"
                n.dismiss()
                rerender()
            except Exception as e:
                n.dismiss()
                ui.notify(f"Classification failed: {e}", type="negative", timeout=8000)

        ui.button("🧠 Classify with AI", on_click=classify_now).props("color=primary size=lg").classes("w-full mt-4")


def _render_report_result(state: dict, rerender):
    cls = state["classification"]
    is_reportable = cls.get("is_reportable", False)

    with ui.card().classes(f"w-full p-6 {'border-l-4 border-red-500' if is_reportable else 'border-l-4 border-green-500'}"):
        with ui.row().classes("items-start gap-4 w-full"):
            ui.icon("warning" if is_reportable else "check_circle", size="3rem").classes("text-red-600" if is_reportable else "text-green-600")
            with ui.column().classes("gap-1"):
                title = "RIDDOR Reportable" if is_reportable else "Not Reportable Under RIDDOR"
                ui.label(title).classes("text-2xl font-bold")
                if is_reportable:
                    ui.html(f'<span class="badge {TYPE_BADGE.get(cls.get("category"), "badge-gray")}">{cls.get("category_label")}</span>')
                ui.label(f"Confidence: {cls.get('confidence', 'unknown').upper()}").classes("text-sm text-slate-500")

        ui.label(cls.get("reasoning", "")).classes("text-slate-700 mt-3 leading-relaxed")

    if is_reportable:
        with ui.grid(columns=3).classes("w-full gap-4"):
            with ui.card().classes("p-4"):
                ui.label("⏰ Reporting Deadline").classes("font-semibold")
                ui.label(cls.get("reporting_deadline", "N/A")).classes("text-sm text-slate-700")
            with ui.card().classes("p-4"):
                ui.label("🌐 Reporting Method").classes("font-semibold")
                ui.label(cls.get("reporting_method", "N/A")).classes("text-sm text-slate-700")
            with ui.card().classes("p-4"):
                ui.label("📋 Records to Keep").classes("font-semibold")
                for r in cls.get("records_to_keep", []):
                    ui.label(f"• {r}").classes("text-sm text-slate-700")

    with ui.grid(columns=2).classes("w-full gap-4"):
        with ui.card().classes("p-4"):
            ui.label("Key Factors").classes("font-semibold")
            for f in cls.get("key_factors", []):
                ui.label(f"🔹 {f}").classes("text-sm text-slate-700")
        with ui.card().classes("p-4"):
            ui.label("Required Actions").classes("font-semibold")
            for a in cls.get("actions_required", []):
                ui.label(f"✅ {a}").classes("text-sm text-slate-700")

    def save_case():
        form = state.get("form", {})
        category = cls.get("category", "unknown")
        deadline = None
        try:
            d = datetime.strptime(form.get("incident_date") or date.today().isoformat(), "%Y-%m-%d")
            if category in ("death", "specified_injury", "occupational_disease", "dangerous_occurrence", "non_worker_hospital"):
                deadline = (date.today() + timedelta(days=10)).isoformat()
            elif category == "over_7_day":
                deadline = (d + timedelta(days=15)).strftime("%Y-%m-%d")
        except ValueError:
            pass

        new_id = f"inc-{uuid.uuid4().hex[:6]}"
        ref = f"RIDDOR-{date.today().strftime('%Y%m%d')}-{new_id[-4:].upper()}"
        new_inc = {
            "id": new_id, "reference": ref,
            "created_at": datetime.utcnow().isoformat(),
            "incident_date": form.get("incident_date") or date.today().isoformat(),
            "incident_type": category,
            "person_name": form.get("person_name") or "Unknown",
            "person_type": form.get("person_type") or "worker",
            "department": form.get("department") or "Unknown",
            "location": form.get("location") or "",
            "description": form.get("description") or "",
            "injury_details": form.get("injury_details") or "",
            "ai_classification": cls,
            "ai_reasoning": cls.get("reasoning") or "",
            "manager_override": None, "status": "open",
            "reporting_deadline": deadline,
            "hse_reference": None, "submitted_at": None,
            "reporter_name": form.get("reporter_name") or "",
            "absence_days": None,
        }
        incidents[new_id] = new_inc
        actions.append({
            "id": uuid.uuid4().hex[:6], "incident_id": new_id, "action_type": "created",
            "description": "Incident report created with AI classification",
            "performed_by": form.get("reporter_name") or "System",
            "timestamp": datetime.utcnow().isoformat(),
        })
        state["new_ref"] = ref
        state["step"] = "saved"
        rerender()

    with ui.row().classes("gap-2 mt-4"):
        ui.button("✅ Accept & Create Case", on_click=save_case).props("color=primary size=lg")
        ui.button("← Go Back & Edit", on_click=lambda: (_back(state, rerender))).props("flat size=lg")


def _back(state: dict, rerender):
    state["step"] = "input"
    state["classification"] = None
    rerender()


def _render_report_saved(state: dict):
    with ui.card().classes("w-full p-8 items-center"):
        ui.icon("check_circle", size="3rem").classes("text-green-500")
        ui.label(f"Case Created: {state.get('new_ref', '')}").classes("text-2xl font-bold mt-2")
        ui.label("The incident is now visible in the Case Monitor.").classes("text-slate-600")
        with ui.row().classes("gap-2 mt-4"):
            ui.button("View Case Monitor", on_click=lambda: ui.navigate.to("/riddor")).props("color=primary")
            ui.button("Report another incident", on_click=lambda: ui.navigate.to("/riddor/report")).props("flat")


# ── RIDDOR: Past Reports (Archive) ───────────────────────────────────

@ui.page("/riddor/archive")
def archive_page():
    app_header(active="riddor")
    with page_container():
        section_header("folder", "Past Reports", "Historical RIDDOR submissions and trends")
        riddor_subnav("archive")

        all_inc = list(incidents.values())
        if not all_inc:
            ui.label("No incidents recorded yet.").classes("text-slate-500")
            return

        # Aggregate counts for the charts
        type_counts: dict = {}
        for i in all_inc:
            label = INCIDENT_TYPE_CONFIG.get(i["incident_type"], {}).get("label", i["incident_type"])
            type_counts[label] = type_counts.get(label, 0) + 1
        dept_counts: dict = {}
        for i in all_inc:
            dept_counts[i["department"]] = dept_counts.get(i["department"], 0) + 1

        # Plotly charts
        try:
            import plotly.graph_objects as go
            with ui.grid(columns=2).classes("w-full gap-4"):
                fig1 = go.Figure(data=[go.Pie(labels=list(type_counts.keys()), values=list(type_counts.values()), hole=0.4)])
                fig1.update_layout(title="Incidents by Type", height=320, margin=dict(t=40, b=20))
                ui.plotly(fig1).classes("w-full")
                fig2 = go.Figure(data=[go.Bar(
                    y=list(dept_counts.keys()), x=list(dept_counts.values()), orientation="h",
                    marker=dict(color="#2563eb"),
                )])
                fig2.update_layout(title="Incidents by Department", height=320, margin=dict(t=40, b=20))
                ui.plotly(fig2).classes("w-full")
        except Exception:
            pass

        # Filters
        search_input = ui.input(placeholder="🔍 Search by name, reference, or description").props("outlined dense").classes("w-full")
        type_filter = ui.select(
            options=["All"] + [v["label"] for v in INCIDENT_TYPE_CONFIG.values()],
            value="All", label="Incident Type",
        ).props("outlined dense")
        dept_filter = ui.select(
            options=["All"] + sorted({i["department"] for i in all_inc}),
            value="All", label="Department",
        ).props("outlined dense")

        results_container = ui.column().classes("w-full gap-2")

        def render_results():
            results_container.clear()
            search = (search_input.value or "").lower()
            t_f = type_filter.value
            d_f = dept_filter.value
            filtered = []
            for i in all_inc:
                if search and not (
                    search in i["person_name"].lower()
                    or search in i["description"].lower()
                    or search in i["reference"].lower()
                ):
                    continue
                if t_f != "All" and INCIDENT_TYPE_CONFIG.get(i["incident_type"], {}).get("label") != t_f:
                    continue
                if d_f != "All" and i["department"] != d_f:
                    continue
                filtered.append(i)
            filtered.sort(key=lambda x: x["incident_date"], reverse=True)
            with results_container:
                ui.label(f"Showing {len(filtered)} of {len(all_inc)} incidents").classes("text-sm text-slate-500")
                for inc in filtered:
                    _render_archive_row(inc)

        search_input.on("update:model-value", lambda e: render_results())
        type_filter.on("update:model-value", lambda e: render_results())
        dept_filter.on("update:model-value", lambda e: render_results())
        render_results()


def _render_archive_row(inc: dict):
    with ui.card().classes("w-full p-4"):
        with ui.row().classes("items-center w-full gap-4 flex-wrap"):
            with ui.column().classes("gap-0"):
                ui.label(inc["reference"]).classes("font-mono font-semibold")
                ui.label(inc["incident_date"]).classes("text-xs text-slate-500")
            ui.html(_type_badge(inc["incident_type"]))
            with ui.column().classes("gap-0"):
                ui.label(inc["person_name"]).classes("font-semibold")
                ui.label(inc["department"]).classes("text-xs text-slate-500")
            ui.space()
            ui.html(_status_badge(inc["status"]))
            if inc.get("hse_reference"):
                ui.label(f"HSE: {inc['hse_reference']}").classes("text-xs font-mono text-slate-500")

        with ui.expansion("View full report").classes("w-full"):
            ui.label("Description").classes("font-semibold mt-2")
            ui.label(inc["description"]).classes("text-slate-700")
            if inc.get("injury_details"):
                ui.label("Injury Details").classes("font-semibold mt-2")
                ui.label(inc["injury_details"]).classes("text-slate-700")
            if inc.get("ai_reasoning"):
                ui.label("AI Assessment").classes("font-semibold mt-2")
                ui.label(inc["ai_reasoning"]).classes("text-slate-700")


# ── RIDDOR: AI Assistant ─────────────────────────────────────────────

@ui.page("/riddor/assistant")
def assistant_page():
    app_header(active="riddor")
    with page_container():
        section_header("chat", "AI Assistant", "Ask about RIDDOR rules or your own data")
        riddor_subnav("assistant")

        chat_box = ui.column().classes("w-full gap-3")

        def render_messages():
            chat_box.clear()
            with chat_box:
                if not ai_chat_messages:
                    with ui.column().classes("items-center w-full py-8 gap-2"):
                        ui.icon("shield", size="3rem").classes("text-blue-500")
                        ui.label("Ask anything about RIDDOR or your data").classes("text-lg font-semibold")
                        ui.label("Try the prompts below or type your own question.").classes("text-sm text-slate-500")
                    with ui.grid(columns=2).classes("w-full gap-2 mt-4"):
                        for q in [
                            "Is a broken finger RIDDOR reportable?",
                            "Which cases are overdue and by how much?",
                            "What counts as a dangerous occurrence?",
                            "Tell me about the kitchen burns case",
                            "How do I calculate the 7-day absence period?",
                            "How many incidents has each department had?",
                        ]:
                            ui.button(q, on_click=lambda q=q: send(q)).props("flat outline").classes("text-left")
                for m in ai_chat_messages:
                    if m["role"] == "user":
                        with ui.row().classes("w-full justify-end"):
                            with ui.element("div").classes("max-w-3xl px-4 py-3 rounded-2xl bg-blue-600 text-white"):
                                ui.label(m["content"])
                    else:
                        with ui.row().classes("w-full"):
                            with ui.element("div").classes("max-w-3xl px-4 py-3 rounded-2xl bg-slate-100 text-slate-900"):
                                ui.markdown(m["content"])

        def send(message: str):
            if not message.strip():
                return
            ai_chat_messages.append({"role": "user", "content": message})
            render_messages()
            n = ui.notification(message="Thinking...", spinner=True, timeout=None)
            try:
                from riddor_ai import data_chat
                sds_dict = {sds_doc["name"]: {"structured": sds_doc["structured"]}} if sds_doc else {}
                ans = data_chat(ai_chat_messages, list(incidents.values()), actions, sds_dict)
                ai_chat_messages.append({"role": "assistant", "content": ans})
            except Exception as e:
                ai_chat_messages.append({"role": "assistant", "content": f"Sorry, I encountered an error: {e}"})
            n.dismiss()
            render_messages()

        render_messages()
        with ui.row().classes("w-full mt-2"):
            input_box = ui.input(placeholder="Ask about RIDDOR or your data...").props("outlined dense").classes("flex-1")

            def on_send():
                msg = input_box.value or ""
                input_box.value = ""
                send(msg)

            input_box.on("keydown.enter", lambda e: on_send())
            ui.button("Send", on_click=on_send).props("color=primary")
            if ai_chat_messages:
                def clear_chat():
                    ai_chat_messages.clear()
                    render_messages()
                ui.button("🗑️", on_click=clear_chat).props("flat dense")


# ── COSHH ────────────────────────────────────────────────────────────

@ui.page("/coshh")
def coshh_page():
    app_header(active="coshh")
    with page_container():
        section_header("science", "COSHH Assistant", "Upload a Safety Data Sheet, generate a condensed assessment, fill your template")

        body = ui.column().classes("w-full gap-4")

        def render():
            body.clear()
            with body:
                _render_coshh(render)

        render()


def _render_coshh(rerender):
    global sds_doc

    # Upload section
    with ui.card().classes("w-full p-4"):
        ui.label("📄 Upload a Safety Data Sheet").classes("text-lg font-semibold mb-1")
        ui.label("PDF, DOCX or TXT").classes("text-xs text-slate-500")

        async def handle_upload(e: events.UploadEventArguments):
            global sds_doc
            data = e.content.read()
            n = ui.notification(message=f"Reading {e.name}...", spinner=True, timeout=None)
            try:
                text = _extract_doc_text(e.name, data)
                if not text.strip():
                    n.dismiss()
                    ui.notify("Could not extract any text from this file", type="negative")
                    return
                n.message = "AI is generating the COSHH assessment..."
                from riddor_ai import extract_sds
                structured = extract_sds(text)
                sds_doc = {"name": e.name, "text": text, "structured": structured}
                global coshh_chat_messages
                coshh_chat_messages = []
                n.dismiss()
                ui.notify(f"Assessment generated for {e.name}", type="positive")
                rerender()
            except Exception as ex:
                n.dismiss()
                ui.notify(f"Failed: {ex}", type="negative", timeout=8000)

        ui.upload(on_upload=handle_upload, auto_upload=True, max_file_size=20_000_000).props("accept=.pdf,.docx,.txt").classes("w-full")

    if sds_doc is None:
        return

    structured = sds_doc["structured"]
    if "_raw_response" in structured:
        with ui.card().classes("w-full p-4 border-l-4 border-amber-500"):
            ui.label("Extraction failed — raw model output:").classes("font-semibold")
            ui.label(structured.get("_raw_response", "")).classes("text-xs text-slate-600 whitespace-pre-wrap")
        return

    _render_assessment(structured)
    _render_template_section(structured, sds_doc["name"])
    _render_coshh_chat(rerender)

    ui.button("📄 New document", on_click=lambda: _coshh_reset(rerender)).props("flat").classes("mt-2")


def _coshh_reset(rerender):
    global sds_doc, coshh_chat_messages
    sds_doc = None
    coshh_chat_messages = []
    rerender()


def _render_assessment(structured: dict):
    product = structured.get("product") or {}
    hazards = structured.get("hazards") or {}
    ppe = structured.get("ppe") or {}
    first_aid = structured.get("first_aid") or {}
    storage = structured.get("storage") or {}

    with ui.card().classes("w-full p-6"):
        ui.label("📋 Condensed COSHH Assessment").classes("text-xl font-bold mb-3")

        with ui.row().classes("w-full justify-between items-start"):
            with ui.column().classes("gap-0"):
                ui.label(product.get("name") or "Unknown product").classes("text-2xl font-bold")
                bits = []
                if product.get("code"): bits.append(f"Code: {product['code']}")
                if product.get("cas_number"): bits.append(f"CAS: {product['cas_number']}")
                if structured.get("physical_state"): bits.append(structured["physical_state"])
                if bits:
                    ui.label(" · ".join(bits)).classes("text-sm text-slate-500")
            signal = (hazards.get("signal_word") or "").upper()
            if signal:
                color = "#dc2626" if signal == "DANGER" else "#d97706"
                ui.html(f'<span style="background:{color};color:#fff;padding:6px 16px;border-radius:6px;font-weight:700">{signal}</span>')

        if structured.get("summary"):
            with ui.element("div").classes("bg-blue-50 text-blue-900 p-3 rounded-lg my-3"):
                ui.label(structured["summary"])

        def field_row(label: str, value):
            if not value:
                return
            with ui.row().classes("gap-2 items-baseline mb-1"):
                ui.label(f"{label}:").classes("font-bold").style("min-width: 200px")
                ui.label(str(value)).classes("text-slate-700")

        def section(title: str, body):
            if not body:
                return
            ui.label(f"{title}:").classes("text-blue-600 font-bold mt-3")
            if isinstance(body, list):
                for item in body:
                    if item:
                        ui.markdown(f"- {item}").classes("ml-2")
            elif isinstance(body, dict):
                for k, v in body.items():
                    if v:
                        ui.markdown(f"- **{k}:** {v}").classes("ml-2")
            else:
                ui.label(str(body)).classes("text-slate-700")

        field_row("Use of Product", structured.get("use_of_product"))
        field_row("Form", structured.get("form") or structured.get("appearance"))
        field_row("pH", structured.get("ph"))

        section("Method of Application", structured.get("method_of_application"))

        clp = structured.get("clp_hazard_summary") or ""
        if hazards.get("h_statements"):
            clp = (clp + "\n\n" + "\n".join(f"- {h}" for h in hazards["h_statements"])).strip()
        section("CLP Hazard Information", clp)

        if structured.get("routes_of_entry"):
            section("Routes of Entry", ", ".join(structured["routes_of_entry"]))

        ppe_items = [f"**{k}:** {v}" for k, v in [
            ("Hands", ppe.get("hands")), ("Eyes", ppe.get("eyes")),
            ("Respiratory", ppe.get("respiratory")), ("Body", ppe.get("body")),
        ] if v]
        section("Mandatory PPE", ppe_items)

        fa_dict = {}
        for k, label in [
            ("skin_contact", "Skin contact"), ("eye_contact", "Eye contact"),
            ("inhalation", "Inhalation"), ("ingestion", "Ingestion"),
        ]:
            if first_aid.get(k):
                fa_dict[label] = first_aid[k]
        section("First Aid Measures", fa_dict)

        section("Spillage Procedure", structured.get("spill_response"))
        storage_parts = []
        if storage.get("conditions"): storage_parts.append(storage["conditions"])
        if storage.get("container"): storage_parts.append(f"Container: {storage['container']}")
        if storage.get("incompatible_materials"): storage_parts.append(f"Keep away from: {storage['incompatible_materials']}")
        section("Handling and Storage", " ".join(storage_parts) if storage_parts else None)
        section("General Precautions", structured.get("general_precautions"))
        section("Disposal", structured.get("disposal"))

        rating = (structured.get("risk_rating") or "LOW").upper()
        rating_color = "#16a34a" if "LOW" in rating else "#dc2626" if "HIGH" in rating else "#d97706"
        with ui.row().classes("items-center mt-4 gap-2"):
            ui.label("Risk Rating (after control measures):").classes("font-bold")
            ui.html(f'<span style="background:{rating_color};color:#fff;padding:4px 14px;border-radius:6px;font-weight:700">{rating}</span>')


def _render_template_section(structured: dict, source_name: str):
    """Template selector + uploader + fill download."""
    from template_store import list_templates, save_template, get_template_bytes, delete_template, BUILTIN_KEY

    with ui.card().classes("w-full p-4"):
        ui.label("📥 Fill COSHH Template").classes("text-lg font-semibold mb-2")
        ui.label("Pick a stored template, or upload a new one — uploaded templates are saved in the list.").classes("text-xs text-slate-500 mb-3")

        templates_list = list_templates()
        options = {key: name for key, name in templates_list}

        global selected_template_key
        if selected_template_key not in options:
            selected_template_key = BUILTIN_KEY

        with ui.row().classes("w-full gap-2 items-center"):
            sel = ui.select(options, value=selected_template_key, label="Template").props("outlined dense").classes("flex-1")
            def on_select(e):
                global selected_template_key
                selected_template_key = sel.value
            sel.on("update:model-value", on_select)

            def remove_selected():
                if selected_template_key == BUILTIN_KEY:
                    ui.notify("Built-in template can't be removed", type="warning")
                    return
                delete_template(selected_template_key)
                ui.notify("Template removed", type="positive")
                ui.navigate.reload()

            ui.button("🗑️", on_click=remove_selected).props("flat dense").tooltip("Remove selected template")

        # Upload new template
        with ui.row().classes("w-full mt-2 items-center gap-2"):
            ui.label("Upload new template:").classes("text-sm text-slate-600")

            async def on_template_upload(e: events.UploadEventArguments):
                if not e.name.lower().endswith(".docx"):
                    ui.notify("Templates must be .docx", type="warning")
                    return
                key = save_template(e.name, e.content.read())
                ui.notify(f"Template saved: {key}", type="positive")
                global selected_template_key
                selected_template_key = key
                ui.navigate.reload()

            ui.upload(on_upload=on_template_upload, auto_upload=True, max_file_size=10_000_000).props("accept=.docx flat dense").classes("flex-1")

        with ui.row().classes("w-full mt-3 gap-2"):
            ui.button("📋 Show placeholders", on_click=_show_placeholders_dialog).props("flat")

            async def fill_and_download():
                from coshh_docx import fill_template
                try:
                    template_bytes = get_template_bytes(selected_template_key)
                    filled = fill_template(template_bytes, structured)
                    product_name = (structured.get("product") or {}).get("name") or source_name.rsplit(".", 1)[0]
                    safe = "".join(c if c.isalnum() else "_" for c in product_name)[:50]
                    fname = f"COSHH_{safe}.docx"
                    ui.download(filled, fname)
                    ui.notify("Filled template downloaded", type="positive")
                except Exception as ex:
                    ui.notify(f"Fill failed: {ex}", type="negative", timeout=8000)

            ui.button("⬇️ Generate & Download", on_click=fill_and_download).props("color=primary")


def _show_placeholders_dialog():
    from coshh_docx import PLACEHOLDERS
    with ui.dialog() as dlg, ui.card().classes("w-full max-w-xl"):
        ui.label("Template Placeholders").classes("text-lg font-bold mb-2")
        ui.label("Add these to your custom template (curly braces matter):").classes("text-sm text-slate-600 mb-3")
        with ui.column().classes("gap-2"):
            for key, desc in PLACEHOLDERS:
                with ui.row().classes("items-center gap-3"):
                    ui.code(f"{{{{{key}}}}}").classes("font-mono")
                    ui.label(desc).classes("text-sm text-slate-600")
        ui.button("Close", on_click=dlg.close).classes("mt-3")
    dlg.open()


def _render_coshh_chat(rerender):
    with ui.card().classes("w-full p-4"):
        ui.label(f"💬 Ask questions about {sds_doc['name']}").classes("text-lg font-semibold mb-2")

        chat_container = ui.column().classes("w-full gap-2")

        def show_messages():
            chat_container.clear()
            with chat_container:
                if not coshh_chat_messages:
                    with ui.grid(columns=2).classes("w-full gap-2"):
                        for q in ["What PPE do I need?", "First aid for eye contact?", "How should this be stored?", "Is this flammable?"]:
                            ui.button(q, on_click=lambda q=q: send(q)).props("flat outline")
                for m in coshh_chat_messages:
                    if m["role"] == "user":
                        with ui.row().classes("w-full justify-end"):
                            with ui.element("div").classes("max-w-3xl px-4 py-3 rounded-2xl bg-blue-600 text-white"):
                                ui.label(m["content"])
                    else:
                        with ui.row().classes("w-full"):
                            with ui.element("div").classes("max-w-3xl px-4 py-3 rounded-2xl bg-slate-100 text-slate-900"):
                                ui.markdown(m["content"])

        def send(message: str):
            if not message.strip():
                return
            coshh_chat_messages.append({"role": "user", "content": message})
            show_messages()
            n = ui.notification(message="Reading the document...", spinner=True, timeout=None)
            try:
                from riddor_ai import sds_chat
                ans = sds_chat(coshh_chat_messages, sds_doc["text"], sds_doc["structured"])
                coshh_chat_messages.append({"role": "assistant", "content": ans})
            except Exception as e:
                coshh_chat_messages.append({"role": "assistant", "content": f"Error: {e}"})
            n.dismiss()
            show_messages()

        show_messages()
        with ui.row().classes("w-full mt-2"):
            inp = ui.input(placeholder=f"Ask about {sds_doc['name']}...").props("outlined dense").classes("flex-1")
            def on_enter():
                msg = inp.value or ""
                inp.value = ""
                send(msg)
            inp.on("keydown.enter", lambda e: on_enter())
            ui.button("Send", on_click=on_enter).props("color=primary")


def _extract_doc_text(name: str, data: bytes) -> str:
    n = name.lower()
    if n.endswith(".pdf"):
        from pypdf import PdfReader
        return "\n\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(data)).pages)
    if n.endswith(".docx"):
        from docx import Document
        d = Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if n.endswith(".txt"):
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {name}")


# ── Run ──────────────────────────────────────────────────────────────

if __name__ in {"__main__", "__mp_main__"}:
    ui.run(
        host="0.0.0.0",
        port=int(os.environ.get("PORT", 8000)),
        title="H&S Hub",
        reload=False,
        show=False,
    )
