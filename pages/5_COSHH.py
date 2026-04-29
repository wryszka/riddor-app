"""COSHH Safety Data Sheet upload, extraction and Q&A."""

import io
import streamlit as st

st.markdown("## 🧪 COSHH Assistant")
st.caption("Upload chemical Safety Data Sheets to extract key COSHH information and ask questions")

# ── Session state ─────────────────────────────────────────────────────
if "sds_documents" not in st.session_state:
    # Map: filename -> {"text": str, "structured": dict, "chat": list}
    st.session_state.sds_documents = {}
if "active_sds" not in st.session_state:
    st.session_state.active_sds = None


# ── Text extraction helpers ──────────────────────────────────────────

def extract_pdf_text(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n\n".join((p.extract_text() or "") for p in reader.pages)


def extract_docx_text(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    if name.endswith(".pdf"):
        return extract_pdf_text(data)
    if name.endswith(".docx"):
        return extract_docx_text(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {name}")


# ── Upload section ────────────────────────────────────────────────────

uploaded = st.file_uploader(
    "Upload a Safety Data Sheet (PDF, DOCX, or TXT)",
    type=["pdf", "docx", "txt"],
    accept_multiple_files=False,
    key="sds_uploader",
)

if uploaded and uploaded.name not in st.session_state.sds_documents:
    with st.spinner(f"Reading {uploaded.name}..."):
        try:
            text = extract_text(uploaded)
            if not text.strip():
                st.error("Could not extract any text from this file. If it's a scanned PDF, OCR is required.")
            else:
                with st.spinner("AI is extracting COSHH information..."):
                    from riddor_ai import extract_sds
                    structured = extract_sds(text)
                st.session_state.sds_documents[uploaded.name] = {
                    "text": text,
                    "structured": structured,
                    "chat": [],
                }
                st.session_state.active_sds = uploaded.name
                st.toast(f"Extracted {uploaded.name}", icon="✅")
                st.rerun()
        except Exception as e:
            st.error(f"Failed to process file: {e}")


# ── Document selector ────────────────────────────────────────────────

if st.session_state.sds_documents:
    docs = list(st.session_state.sds_documents.keys())
    if st.session_state.active_sds not in docs:
        st.session_state.active_sds = docs[0]

    cols = st.columns([3, 1])
    with cols[0]:
        st.session_state.active_sds = st.selectbox(
            "Active document",
            options=docs,
            index=docs.index(st.session_state.active_sds),
            key="sds_selector",
        )
    with cols[1]:
        st.write("")  # spacer for alignment
        if st.button("🗑️ Remove", use_container_width=True):
            del st.session_state.sds_documents[st.session_state.active_sds]
            st.session_state.active_sds = next(iter(st.session_state.sds_documents), None)
            st.rerun()

    st.markdown("---")

    active = st.session_state.sds_documents[st.session_state.active_sds]
    structured = active["structured"]

    # ── Tabs: Summary | Q&A | Raw text ──────────────────────────────
    tab_summary, tab_chat, tab_raw = st.tabs(["📋 Summary", "💬 Ask Questions", "📄 Raw Text"])

    with tab_summary:
        if "_raw_response" in structured:
            st.warning("Extraction failed — showing raw model output below.")
            st.code(structured.get("_raw_response", ""), language=None)
        else:
            # Header
            product = structured.get("product") or {}
            supplier = structured.get("supplier") or {}
            hazards = structured.get("hazards") or {}

            name = product.get("name") or "Unknown product"
            st.markdown(f"### {name}")
            if structured.get("summary"):
                st.info(structured["summary"])

            # Top-line facts in columns
            c1, c2, c3 = st.columns(3)
            with c1:
                st.markdown("**Product**")
                if product.get("code"): st.caption(f"Code: {product['code']}")
                if product.get("cas_number"): st.caption(f"CAS: {product['cas_number']}")
                if product.get("ec_number"): st.caption(f"EC: {product['ec_number']}")
                if structured.get("physical_state"): st.caption(f"State: {structured['physical_state']}")
                if structured.get("appearance"): st.caption(f"Appearance: {structured['appearance']}")
            with c2:
                st.markdown("**Supplier**")
                if supplier.get("name"): st.caption(supplier["name"])
                if supplier.get("phone"): st.caption(f"📞 {supplier['phone']}")
                if supplier.get("emergency_phone"): st.caption(f"🚨 Emergency: {supplier['emergency_phone']}")
            with c3:
                st.markdown("**Signal Word**")
                signal = hazards.get("signal_word")
                if signal:
                    color = "#dc2626" if signal.lower() == "danger" else "#d97706" if signal.lower() == "warning" else "#6b7280"
                    st.markdown(f'<span style="background:{color};color:white;padding:4px 12px;border-radius:6px;font-weight:600">{signal.upper()}</span>', unsafe_allow_html=True)
                if hazards.get("pictograms"):
                    st.caption("Pictograms:")
                    for p in hazards["pictograms"]:
                        st.caption(f"• {p}")

            st.markdown("---")

            # Hazards section
            st.markdown("#### ⚠️ Hazards")
            ch1, ch2 = st.columns(2)
            with ch1:
                st.markdown("**GHS Classifications**")
                for c in (hazards.get("ghs_classifications") or []):
                    st.markdown(f"- {c}")
                if not hazards.get("ghs_classifications"):
                    st.caption("Not specified")
            with ch2:
                st.markdown("**Hazard (H) Statements**")
                for h in (hazards.get("h_statements") or []):
                    st.markdown(f"- {h}")
                if not hazards.get("h_statements"):
                    st.caption("Not specified")

            st.markdown("**Precautionary (P) Statements**")
            p_statements = hazards.get("p_statements") or []
            if p_statements:
                with st.expander(f"View {len(p_statements)} precautionary statements"):
                    for p in p_statements:
                        st.markdown(f"- {p}")
            else:
                st.caption("Not specified")

            st.markdown("---")

            # PPE
            st.markdown("#### 🥽 Personal Protective Equipment")
            ppe = structured.get("ppe") or {}
            ppe_cols = st.columns(4)
            ppe_items = [
                ("Eyes/Face", ppe.get("eyes")),
                ("Hands", ppe.get("hands")),
                ("Respiratory", ppe.get("respiratory")),
                ("Body", ppe.get("body")),
            ]
            for col, (label, value) in zip(ppe_cols, ppe_items):
                with col:
                    st.markdown(f"**{label}**")
                    st.caption(value or "Not specified")

            st.markdown("---")

            # First aid
            st.markdown("#### 🩹 First Aid Measures")
            fa = structured.get("first_aid") or {}
            fa_cols = st.columns(2)
            fa_items = [
                ("👁️ Eye Contact", fa.get("eye_contact")),
                ("🤲 Skin Contact", fa.get("skin_contact")),
                ("🫁 Inhalation", fa.get("inhalation")),
                ("🍽️ Ingestion", fa.get("ingestion")),
            ]
            for i, (label, value) in enumerate(fa_items):
                with fa_cols[i % 2]:
                    st.markdown(f"**{label}**")
                    st.caption(value or "Not specified")

            st.markdown("---")

            # Storage and exposure limits
            sc1, sc2 = st.columns(2)
            with sc1:
                st.markdown("#### 📦 Storage")
                storage = structured.get("storage") or {}
                if storage.get("conditions"): st.markdown(f"**Conditions:** {storage['conditions']}")
                if storage.get("container"): st.markdown(f"**Container:** {storage['container']}")
                if storage.get("incompatible_materials"): st.markdown(f"**Keep away from:** {storage['incompatible_materials']}")
                if not any(storage.values() if storage else []):
                    st.caption("Not specified")
            with sc2:
                st.markdown("#### 📊 Exposure Limits")
                limits = structured.get("exposure_limits") or {}
                if limits.get("wel_8h_twa"): st.markdown(f"**8h WEL:** {limits['wel_8h_twa']}")
                if limits.get("stel_15min"): st.markdown(f"**15min STEL:** {limits['stel_15min']}")
                if limits.get("biological_limits"): st.markdown(f"**Biological:** {limits['biological_limits']}")
                if not any(limits.values() if limits else []):
                    st.caption("Not specified")

            st.markdown("---")

            # Fire, spill, disposal
            with st.expander("🔥 Fire-fighting measures"):
                ff = structured.get("fire_fighting") or {}
                if ff.get("extinguishing_media"): st.markdown(f"**Suitable:** {ff['extinguishing_media']}")
                if ff.get("unsuitable_media"): st.markdown(f"**Unsuitable:** {ff['unsuitable_media']}")
                if ff.get("hazardous_combustion"): st.markdown(f"**Hazardous combustion products:** {ff['hazardous_combustion']}")

            with st.expander("💧 Spill response"):
                st.markdown(structured.get("spill_response") or "Not specified")

            with st.expander("♻️ Disposal"):
                st.markdown(structured.get("disposal") or "Not specified")

    with tab_chat:
        st.caption(f"Asking questions about **{st.session_state.active_sds}**")

        # Example prompts when chat empty
        if not active["chat"]:
            EXAMPLES = [
                "What PPE do I need to use this safely?",
                "What's the first aid procedure if it gets in someone's eye?",
                "Is this substance flammable?",
                "What's the WEL exposure limit?",
                "How should this be stored?",
                "What should I do if it spills?",
            ]
            st.markdown("**Try asking:**")
            ec = st.columns(2)
            for i, q in enumerate(EXAMPLES):
                if ec[i % 2].button(q, key=f"sds_ex_{i}", use_container_width=True):
                    active["chat"].append({"role": "user", "content": q})
                    with st.spinner("Reading the document..."):
                        try:
                            from riddor_ai import sds_chat
                            ans = sds_chat(active["chat"], active["text"], active["structured"])
                            active["chat"].append({"role": "assistant", "content": ans})
                        except Exception as e:
                            active["chat"].append({"role": "assistant", "content": f"Error: {e}"})
                    st.rerun()

        # Display chat
        for msg in active["chat"]:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        # Input
        if prompt := st.chat_input(f"Ask about {st.session_state.active_sds}..."):
            active["chat"].append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            with st.chat_message("assistant"):
                with st.spinner("Reading the document..."):
                    try:
                        from riddor_ai import sds_chat
                        ans = sds_chat(active["chat"], active["text"], active["structured"])
                        st.markdown(ans)
                        active["chat"].append({"role": "assistant", "content": ans})
                    except Exception as e:
                        msg = f"Error: {e}"
                        st.markdown(msg)
                        active["chat"].append({"role": "assistant", "content": msg})

        if active["chat"]:
            if st.button("🗑️ Clear conversation", key="sds_clear_chat"):
                active["chat"] = []
                st.rerun()

    with tab_raw:
        st.caption("Raw text extracted from the document — for verification")
        st.text_area(
            "Document text",
            value=active["text"],
            height=500,
            label_visibility="collapsed",
            key=f"raw_text_{st.session_state.active_sds}",
        )

else:
    st.info("👆 Upload a Safety Data Sheet (SDS) to get started. Supported formats: PDF, DOCX, TXT.")

st.caption("⚠️ AI-assisted summary — always verify against the original SDS before making safety decisions")
